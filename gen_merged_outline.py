from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ==================== 全局样式设置 ====================
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)

def add_title(text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x5C, 0x6E)
    return heading

def add_paragraph_text(text, bold=False, size=10.5):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_styled_table(headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '0A5C6E')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(10)

    # 数据行
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            if r % 2 == 1:
                set_cell_shading(cell, 'F5F5F0')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)

    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

# ==================== 文档正文 ====================

# 大标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【生命健康始于肠道】合并版PPT大纲')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x0A, 0x5C, 0x6E)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A版（企业权威型）× B版（认知引导型）最优合并 · 30页高净值客户产说会')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('合并策略：B版开场（让客户想买）→ A版中段（让客户敢买）→ B版收口（家庭方案成交）')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xB0, 0x75, 0x17)
run.italic = True

doc.add_paragraph()

# ==================== 一、整体参数 ====================
add_title('一、整体参数', 2)

add_styled_table(
    ['参数', '内容'],
    [
        ['总页数', '30页'],
        ['用途', '保险公司高净值客户产说会 / 健康沙龙'],
        ['演讲时长', '40分钟'],
        ['基调', '认知引导 + 专业权威 + 温暖人文（混搭型）'],
        ['开场策略', '家庭健康焦虑唤醒（取自B版）'],
        ['情感驱动', '健康管理趋势焦虑 + 自我识别代入（B版逻辑）'],
        ['信任建立', '企业硬实力 + 精选案例（A版优势）'],
        ['成交模式', '家庭年度微生态健康管理方案（B版收口）'],
        ['合规标准', '全部采用B版合规表达规范'],
    ]
)

doc.add_paragraph()

# ==================== 二、配色方案 ====================
add_title('二、配色方案', 2)

add_styled_table(
    ['角色', '色值', '用途'],
    [
        ['主色', '#0A5C6E', '章节标题、核心信息'],
        ['辅助色', '#2E9C8A', '图标、辅助信息'],
        ['强调色', '#F5A623', '重点标注、限时优惠'],
        ['浅底', '#F5F5F0', '页面背景'],
        ['深底', '#1A2A3A', '深色页面'],
        ['正文', '#333333', '正文文字'],
    ]
)

doc.add_paragraph()

# ==================== 三、章节结构 ====================
add_title('三、章节结构总览', 2)

add_styled_table(
    ['章节', '页码', '页数', '来源', '核心功能'],
    [
        ['封面', 'P01', '1页', 'B版优化', '建立高级感 + 家庭健康守护'],
        ['第一章：家庭健康焦虑唤醒', 'P02-05', '4页', 'B版', '让客户意识到"该管了"'],
        ['第二章：认知强化', 'P06-11', '6页', 'B版精简', '打透肠道子系统，建立科学认知'],
        ['第三章：客户自我识别', 'P12-15', '4页', 'B版', '"说的是不是我？"成交转折点'],
        ['第四章：方案矩阵', 'P16-19', '4页', 'A版精简', '产品详解 + FMT + 对比'],
        ['第五章：企业信任', 'P20-23', '4页', 'A版精简', '硬实力背书，让客户敢买'],
        ['第六章：案例实证', 'P24-26', '3页', 'A版精选', '3个最强案例'],
        ['第七章：行动号召', 'P27-30', '4页', 'B版优化', '家庭方案成交收口'],
    ]
)

doc.add_paragraph()

# ==================== 四、逐页内容 ====================
add_title('四、逐页内容', 2)

# ---------- 封面 ----------
doc.add_paragraph()
p = add_paragraph_text('封面（1页）', bold=True, size=12)
p.runs[0].font.color.rgb = RGBColor(0x0A, 0x5C, 0x6E)

add_styled_table(
    ['页码', '标题', '核心信息', '要点方向', '来源'],
    [
        ['P01', '肠道菌群：高净值家庭的第二健康资产',
         '主标题 + 副标题 + 承葛医药集团 + 保险高净值客户健康沙龙',
         '建立高级感和家庭健康守护感，不要"科普课"感', 'B版'],
    ],
    col_widths=[1.2, 4, 5, 4.5, 1.2]
)

# ---------- 第一章 ----------
doc.add_paragraph()
p = add_paragraph_text('第一章：家庭健康焦虑唤醒（4页）', bold=True, size=12)
p.runs[0].font.color.rgb = RGBColor(0x0A, 0x5C, 0x6E)

add_styled_table(
    ['页码', '标题', '核心信息', '要点方向', '来源'],
    [
        ['P02', '有钱之后，最怕的不是花钱，是健康突然失控',
         '自己、父母、配偶、子女的健康不确定性；财富可解决医疗资源，但不能替代提前管理',
         '唤醒健康风险焦虑，但不制造恐惧', 'B版'],
        ['P03', '高净值家庭真正想守住的，是一家人的长期安心',
         '自己状态稳定、父母慢病风险、配偶共同管理、子女体质与免疫',
         '从个人问题带入家庭健康资产配置', 'B版'],
        ['P04', '为什么很多人体检正常，却感觉状态越来越差？',
         '疲劳、肠胃不适、睡眠变差、体重血糖血脂难管、免疫力下降、应酬后恢复慢',
         '制造认知冲突，引出微生态检测', 'B版'],
        ['P05', '健康管理正在从"查疾病"走向"看趋势"',
         '传统体检：体检→指标异常→医院治疗；高端管理：早期评估→趋势观察→精准干预→复测维护',
         '提升客户对主动健康管理的认知', 'B版'],
    ],
    col_widths=[1.2, 4, 5, 4.5, 1.2]
)

# ---------- 第二章 ----------
doc.add_paragraph()
p = add_paragraph_text('第二章：认知强化（6页）', bold=True, size=12)
p.runs[0].font.color.rgb = RGBColor(0x0A, 0x5C, 0x6E)

add_styled_table(
    ['页码', '标题', '核心信息', '要点方向', '来源'],
    [
        ['P06', '认识你的"第二大脑"：肠道',
         '肠道不只是消化器官，也与神经、情绪、睡眠、压力、食欲和身体状态反馈相关',
         '用"第二大脑"建立强记忆点', 'B版'],
        ['P07', '你不是一个人，你是一个生态系统',
         '人体与大量微生物共生；100万亿菌群、1-3kg总重量；哈佛研究："我们是一群生物的集合体"',
         '形成强认知金句（A版哈佛数据+B版金句）', 'A+B合并'],
        ['P08', '肠道菌群：身体里的隐形管理者',
         '有益菌、条件致病菌、菌群多样性、菌群平衡、代谢产物、微生态稳定性',
         '建立菌群基本概念', 'B版'],
        ['P09', '菌群失衡：许多健康问题背后的底层变量',
         '关联便秘、腹胀、腹泻、口气、睡眠差、皮肤反复、代谢难管等',
         '强化风险感，保持合规（不用"万病之源"）', 'B版合规版'],
        ['P10', '肠脑轴 + 免疫 + 代谢：三个你不知道的关联',
         '压力→肠胃不适；肠道是重要免疫场所；能量利用、血糖血脂管理效率',
         '三合一讲透"为什么肠道影响全身"（B版3页精华压缩）', 'B版精简'],
        ['P11', '肠道屏障：身体健康的第一道防线',
         '肠黏膜、黏液层、免疫防线、菌群生态、肠道通透性',
         '让客户理解肠道不仅是排便', 'B版'],
    ],
    col_widths=[1.2, 4, 5, 4.5, 1.2]
)

# ---------- 第三章 ----------
doc.add_paragraph()
p = add_paragraph_text('第三章：客户自我识别（4页）', bold=True, size=12)
p.runs[0].font.color.rgb = RGBColor(0x0A, 0x5C, 0x6E)

add_styled_table(
    ['页码', '标题', '核心信息', '要点方向', '来源'],
    [
        ['P12', '这些信号，可能正在提醒你关注肠道微生态',
         '便秘、腹胀、排气多、腹泻、口气、消化不良、睡眠差、皮肤反复、体重难管、恢复慢',
         '让客户对号入座（成交关键转折点）', 'B版'],
        ['P13', '为什么你吃了很多保健品，效果却不稳定？',
         '产品吃了不少但改善不稳定；朋友有效自己无感；一停就反复；不知道到底该补什么',
         '制造精准检测需求', 'B版'],
        ['P14', '普通体检看器官，菌群检测看生态',
         '普通体检看器官、血液、影像和异常结果；菌群检测看微生态结构、趋势和平衡',
         '菌群检测是体检补充，不是替代', 'B版'],
        ['P15', '粪便16S检测：精准读懂你的肠道密码',
         '16S rRNA测序：多样性/菌群结构/致病菌筛查/功能预测；无创无痛5分钟采样',
         '具体交付物 + A版技术参数', 'A+B合并'],
    ],
    col_widths=[1.2, 4, 5, 4.5, 1.2]
)

# ---------- 第四章 ----------
doc.add_paragraph()
p = add_paragraph_text('第四章：方案矩阵（4页）', bold=True, size=12)
p.runs[0].font.color.rgb = RGBColor(0x0A, 0x5C, 0x6E)

add_styled_table(
    ['页码', '标题', '核心信息', '要点方向', '来源'],
    [
        ['P16', '肠道管理，不是随便买一盒益生菌',
         '益生菌使用要看菌株、剂量、周期、人群、场景，以及是否与当前菌群状态匹配',
         '建立专业门槛（B版逻辑）', 'B版'],
        ['P17', '承葛益生菌产品矩阵——四款产品，覆盖全家健康需求',
         '健律佳：12株/≥240亿CFU/99%耐胃酸；鑫辰速：三靶点逆糖；盈畅安：修复肠屏障；韵静怡：肠-脑轴双通路',
         '竞品对比 + 差异化优势（A版数据）', 'A版'],
        ['P18', '菌群移植（FMT）——系统重建的临床干预技术',
         '原理：健康人完整菌群生态移植，重建稳态；国家卫健委低风险认证；Science十大突破',
         'FMT原理与合规背景（A版优势）', 'A版'],
        ['P19', '益生菌 × 菌群移植：两种工具，各有舞台',
         '益生菌=补充有益力量，适合日常维护；FMT=重建微生态环境，需专业评估',
         '对比表 + 决策逻辑（A+B双版精华）', 'A+B合并'],
    ],
    col_widths=[1.2, 4, 5, 4.5, 1.2]
)

# ---------- 第五章 ----------
doc.add_paragraph()
p = add_paragraph_text('第五章：企业信任（4页）', bold=True, size=12)
p.runs[0].font.color.rgb = RGBColor(0x0A, 0x5C, 0x6E)

add_styled_table(
    ['页码', '标题', '核心信息', '要点方向', '来源'],
    [
        ['P20', '为什么是承葛？不是每家机构都做得到',
         '10年+研究临床积累；自建超级供体菌库；全流程标准化批批检测可追溯；三甲医院/顶级科研合作',
         '差异化定位 + 行业壁垒', 'A版'],
        ['P21', '承葛的技术底座：10大菌库 × 6大专利技术',
         '10大菌库覆盖全国；CNAS认可；科技部人类遗传资源保藏批件（全国首家）；6大专利技术；供体百里挑一',
         '稀缺资源认知（A版最强页）', 'A版'],
        ['P22', '承葛服务数据：10,000+案例的积累',
         '累计服务超10万例次（行业第一）；适应症30+种；满意度95%+；合作医疗机构100+家',
         '规模背书', 'A版'],
        ['P23', '权威背书：学术 × 监管 × 媒体三重认证',
         '国家卫健委低风险新技术认定；SCI论文发表；沪卫医2021操作规范；中华菌库联盟',
         '多维权威', 'A版'],
    ],
    col_widths=[1.2, 4, 5, 4.5, 1.2]
)

# ---------- 第六章 ----------
doc.add_paragraph()
p = add_paragraph_text('第六章：案例实证（3页）', bold=True, size=12)
p.runs[0].font.color.rgb = RGBColor(0x0A, 0x5C, 0x6E)

add_styled_table(
    ['页码', '标题', '核心信息', '要点方向', '来源'],
    [
        ['P24', '案例一：30年便秘，两疗程改写结局',
         '63岁，便秘30年，伴黑肠病变。两疗程FMT后：1-2次/日规律排便，停用助排药，综合改善',
         '呼应开场的情感故事（如有开场故事则此页为答案揭晓）', 'A版'],
        ['P25', '案例二：血糖13.54→正常，停掉胰岛素',
         '空腹血糖13.54，糖化血红蛋白9.8%。移植后3.4-6.8稳定，"三多一少"消失，停用胰岛素',
         '高净值客户最强共鸣案例（代谢类）', 'A版'],
        ['P26', '案例三：阿尔茨海默症，发表SCI论文的奇迹',
         '76岁额颞叶痴呆，鼻肠管菌液移植，认知改善语言恢复。论文：Frontiers期刊',
         '高难度+学术权威，增强信任感', 'A版'],
    ],
    col_widths=[1.2, 4, 5, 4.5, 1.2]
)

# ---------- 第七章 ----------
doc.add_paragraph()
p = add_paragraph_text('第七章：行动号召（4页）', bold=True, size=12)
p.runs[0].font.color.rgb = RGBColor(0x0A, 0x5C, 0x6E)

add_styled_table(
    ['页码', '标题', '核心信息', '要点方向', '来源'],
    [
        ['P27', '承葛提供的不是单一产品，而是一套微生态管理系统',
         '检测→报告解读→益生菌精准调理→菌群移植方案→个性化干预→管理师跟踪→复测',
         '提高价值感，避免单品价格比较', 'B版'],
        ['P28', '高净值家庭年度微生态健康管理方案',
         '本人/配偶/父母/子女的检测与档案；专家解读；益生菌权益；FMT评估权益；年度复测',
         '从个人购买升级为家庭购买（客单价翻倍）', 'B版'],
        ['P29', '保险客户专属权益',
         '检测名额、专家/管理师解读、个性化建议、家庭加购、益生菌权益、FMT评估权益、复测优惠',
         '现场成交理由', 'B版'],
        ['P30', '今天开始，为家人建立一份肠道微生态健康档案',
         '不是等问题出现，而是提前看见趋势；不是盲目吃产品，而是先检测再管理；不是个人调理，而是家庭资产',
         '温暖成交收口 + 品牌记忆', 'B版'],
    ],
    col_widths=[1.2, 4, 5, 4.5, 1.2]
)

doc.add_paragraph()

# ==================== 五、合并策略说明 ====================
add_title('五、合并策略说明', 2)

add_paragraph_text('1. 为什么P01-05取自B版？')
add_paragraph_text('   B版的开场从"家庭健康资产"切入，比A版的"便秘故事悬念"更贴合高净值客户的身份认同。高净值客户最怕的不是生病，而是"健康失控"。')

add_paragraph_text('2. 为什么P06-11以B版为主？')
add_paragraph_text('   B版的认知逻辑是"跟你什么关系"（肠脑轴影响睡眠、免疫影响全家），比A版的"菌群是什么"（超级生物体、第八大器官）更容易让客户产生代入感。认知阶段的目标不是科普，而是让客户觉得"这事跟我有关"。')

add_paragraph_text('3. 为什么P12-15取自B版？')
add_paragraph_text('   "客户自我识别"是B版独有的设计，A版完全没有这一环节。这5页是成交的关键转折点——让客户从"你们卖什么？"转变为"我是不是也应该查一下？"。')

add_paragraph_text('4. 为什么P16-19以A版为主？')
add_paragraph_text('   讲到产品时，A版的数据更具体（12株/240亿CFU/99%耐胃酸），更有说服力。B版的产品介绍过于笼统。但在P16保留了B版的"不能随便买益生菌"逻辑，作为产品介绍的引子。')

add_paragraph_text('5. 为什么P20-23全取自A版？')
add_paragraph_text('   企业硬实力是A版的核心优势——10大菌库、6大专利、1万+案例、三重认证。B版在这一块几乎空白。这一段让客户"敢买"。')

add_paragraph_text('6. 为什么案例从5个精简到3个？')
add_paragraph_text('   时间有限，保留3个最强案例（便秘30年/糖尿病逆转/阿尔茨海默SCI）。便秘呼应开场故事，糖尿病打中高净值客户代谢痛点，阿尔茨海默展示学术权威。去掉溃疡性结肠炎和荨麻疹/抑郁/失眠。')

add_paragraph_text('7. 为什么P27-30取自B版？')
add_paragraph_text('   B版的收口设计更高级——不是"买盒益生菌"而是"建立家庭年度微生态健康管理方案"，客单价和客户黏性都远高于A版的单品成交模式。')

doc.add_paragraph()

# ==================== 六、合规表达规范 ====================
add_title('六、合规表达规范', 2)

add_paragraph_text('合并版统一采用B版合规标准，所有表述需遵循以下规范：')

add_styled_table(
    ['原表达（禁用）', '风险', '建议改法'],
    [
        ['菌群失衡等于万病之源', '医疗夸大风险高', '菌群失衡是许多健康问题背后的底层变量'],
        ['肠道决定健康的上限', '绝对化', '肠道影响健康管理的底层状态'],
        ['肠道决定一切', '绝对化', '肠道是观察整体健康的重要窗口'],
        ['菌群移植重建健康', '疗效承诺风险', '菌群移植相关技术可用于更专业的微生态干预场景'],
        ['彻底改变/根治/逆转', '疗效承诺风险', '辅助改善、趋势观察、健康管理、个体体验因人而异'],
    ]
)

doc.add_paragraph()

# ==================== 七、讲师节奏建议 ====================
add_title('七、讲师现场节奏建议', 2)

add_styled_table(
    ['阶段', '页码', '节奏建议'],
    [
        ['焦虑唤醒', 'P01-05', '不要急着讲产品。让客户意识到健康风险不是有没有钱的问题，而是能不能提前看见的问题。'],
        ['认知重塑', 'P06-11', '打透3个子系统（肠脑轴/免疫/代谢），让每个客户至少能对上1个自己的症状。'],
        ['自我代入', 'P12-15', '这是成交前最关键的转折。讲好了客户会从"你们卖什么？"转变为"我是不是也应该查一下？"。'],
        ['方案展示', 'P16-19', '先说"不能随便买"（建门槛），再说承葛产品（给方案）。让客户觉得选择你们是"理性的"。'],
        ['信任建立', 'P20-23', '展示硬实力，但不要变成公司介绍会。每页回到"这对您意味着什么"。'],
        ['案例证明', 'P24-26', '讲故事的方式讲案例，不要念数据。便秘案例如有开场呼应则最有效果。'],
        ['成交收口', 'P27-30', '少讲"买产品"，多讲健康档案、家庭年度管理、专属权益和提前看见趋势。'],
    ]
)

doc.add_paragraph()

# ==================== 八、待补充信息清单 ====================
add_title('八、待补充信息清单', 2)

add_styled_table(
    ['优先级', '页码', '待补充字段', '备注'],
    [
        ['高', 'P15', '16S检测报告模板截图', '来自16S模板.pdf'],
        ['中', 'P21', '10大菌库地图/分布图', '可视化素材'],
        ['中', 'P24-26', '案例原始数据/肠镜图/血糖曲线', '合规脱敏后使用'],
        ['高', 'P29', '保险客户专属权益具体条款', '需合作方提供'],
        ['高', 'P29', '现场专属优惠政策', '需市场部确认名额和价格'],
    ]
)

doc.add_paragraph()

# ==================== 尾部金句 ====================
add_title('附录：核心金句', 2)

gold_sentences = [
    '体检看结果，菌群看趋势。',
    '保险解决风险发生后的经济保障，微生态管理关注风险发生前的健康趋势。',
    '你不是一个人，你是一个生态系统。',
    '肠道菌群失衡，可能是许多健康问题反复出现的底层变量。',
    '益生菌像"补充有益力量"，菌群移植更像"重建微生态环境"。',
    '我们今天不是卖一个检测盒，而是帮助每个家庭建立一份肠道微生态健康档案。',
    '真正高端的健康管理，不是等指标异常，而是在身体发出信号前，先看见变化方向。',
]

for s in gold_sentences:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(s)
    run.font.size = Pt(10)
    run.italic = True

# ==================== 保存 ====================
output_path = '/Users/money/WorkBuddy/20260505213642/合并版PPT大纲.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path)} bytes')
